import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from dotenv import load_dotenv
import os

load_dotenv()
# Configurar la API de Gemini
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
model = genai.GenerativeModel('gemini-1.5-flash')





#esta funcion es para extraer el texto relevante y separarlo en comas -------
def extract_sensitive_data_from_pdf(pdf_path, consulta):
    # Cargar el PDF
    doc = fitz.open(pdf_path)

    # Extraer el texto del PDF
    text = ""
    for page in doc:
        text += page.get_text()

    # Imprimir el texto extraído para depuración
    print("Texto extraído del PDF:", text)

    # Generar la respuesta usando la API de Generative AI
    response = model.generate_content(f"{text}\n\nPregunta: {consulta}")
    print(response)

    # Procesar la respuesta y extraer los datos relevantes
    sensitive_data = response._result.candidates[0].content.parts[0].text.split(',')
    sensitive_data = [data.strip() for data in sensitive_data]

    return sensitive_data



#esta funcion extrae el texto de la resolucion y el numero de expediente ------
######
def extract_resolution_and_case_number(pdf_path, consulta):
    # Cargar el PDF
    doc = fitz.open(pdf_path)

    # Extraer el texto del PDF
    text = ""
    for page in doc:
        text += page.get_text()

    # Generar la respuesta usando la API de Generative AI
    response = model.generate_content(f"{text}\n\nPregunta: {consulta}")
    # print(response)

    # Procesar la respuesta y extraer la resolución y el número de expediente
    resolution_data = response._result.candidates[0].content.parts[0].text.split(',')
    resolution_data = [data.strip() for data in resolution_data]

    return resolution_data

# anonimiza los datos relevantes y retorna el pdf anonimizado con BYTESIO
def anonymize_pdf(input_pdf_path, texts_to_anonymize):
    # Abrir el archivo PDF
    document = fitz.open(input_pdf_path)

    # Iterar sobre cada página
    for page_num in range(len(document)):
        page = document[page_num]

        for text in texts_to_anonymize:
            # Buscar el texto a anonimizar
            text_instances = page.search_for(text)

            # Verificar si se encontraron instancias del texto
            if text_instances:
                # Reemplazar cada instancia del texto con un rectángulo negro
                for inst in text_instances:
                    # Añadir una anotación de redacción con un rectángulo negro
                    page.add_redact_annot(inst, fill=(0, 0, 0))

        # Aplicar los cambios de redacción (redaction)
        page.apply_redactions()

    # Guardar el PDF anonimizado en un objeto BytesIO
    output_pdf_stream = BytesIO()
    document.save(output_pdf_stream)
    document.close()

    output_pdf_stream.seek(0)
    return output_pdf_stream


# Función para limpiar el nombre del archivo
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)


def set_font(run, font_name='Calibri', font_size=12):
    """
    Set font size and font name for a run.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def reemplazar_palabra(documento, identificador, texto_deseado):
    for p in documento.paragraphs:
        if identificador in p.text:
            for run in p.runs:
                if identificador in run.text:
                    run.text = run.text.replace(identificador, texto_deseado)
                    set_font(run)

    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if identificador in p.text:
                        for run in p.runs:
                            if identificador in run.text:
                                run.text = run.text.replace(identificador, texto_deseado)
                                set_font(run)


def modify_word_file(resolution, case_number):
    # Determinar el documento a usar basado en la resolución
    print(resolution)
    suffix = resolution[-2:]


    # Determinar el documento a usar basado en los últimos tres caracteres de la resolución
    if suffix == "S1":
        ruta_documento = "SU1.docx"
    elif suffix == "S2":
        ruta_documento = "SU2.docx"
    elif suffix == "SC":
        ruta_documento = "SC.docx"
    else:
        st.error("No se encontró un documento Word correspondiente.")
        return None
    # Abrir el documento existente
    documento = Document(ruta_documento)

    # Reemplazar el identificador "AQUI" con el texto deseado
    identificador_a_reemplazar = "AQUI"
    texto_deseado = f"Se adjunta copia de resolución Nro. {resolution} del expediente Nro. {case_number} conforme se encuentra digitalizada en el sistema electrónico de Osinergmin."
    reemplazar_palabra(documento, identificador_a_reemplazar, texto_deseado)

    # Guardar el documento modificado en un objeto BytesIO
    output_doc_stream = BytesIO()
    documento.save(output_doc_stream)
    output_doc_stream.seek(0)

    return output_doc_stream



# elegir el informe saip atencion parcial
def modify_saip_report(resolution, nombre_solicitante):
    suffix = resolution[-2:]
    if suffix == "S1":
        ruta_documento = "Informe de SAIP_Atención Parcial_S1.docx"
    elif suffix == "S2":
        ruta_documento = "Informe de SAIP_Atención Parcial_S2.docx"
    elif suffix == "SC":
        ruta_documento = "Informe de SAIP_Atención Parcial_SC.docx"
    else:
        st.error("No se encontró un documento Word correspondiente para SAIP atención parcial.")
        return None

    documento = Document(ruta_documento)
    identificador_a_reemplazar = "NAME"
    reemplazar_palabra(documento, identificador_a_reemplazar, nombre_solicitante)

    output_doc_stream = BytesIO()
    documento.save(output_doc_stream)
    output_doc_stream.seek(0)

    return output_doc_stream




# Streamlit App
st.title("Anonimizador de Resol. STOR")

# Incluir estilo CSS personalizado
st.markdown("""
    <style>
    .stTextInput > div > div > input {
        border: 2px solid #4CAF50;
        padding: 5px;
        border-radius: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

# Subir archivo PDF
uploaded_file = st.file_uploader("Sube un archivo PDF", type=["pdf"])

# Añadir un campo de texto para ingresar el nombre del solicitante
nombre_solicitante = st.text_input("Nombre del Solicitante de la SAIP")


if uploaded_file is not None:
    pdf_path = uploaded_file.name
    consulta_datos_sensibles = "dame solamente el nombre del recurrente(s), el suministro(s), la direccion o ubicacion del suministro, la direccion o sitio de notificacion, tercero con interes que no sea la concesionaria, el medidor del suministro, el display del suministro, el numero de serie del medidor, el numero de un nuevo medidor  y correo si los tuviese. cada uno de ellos solo separado por comas y si no ubiese ignorarlas y no incluirlas"
    consulta_resolucion = "dame solamente el numero de la resolución y el número de expediente de la resolución del documento. Cada uno de ellos separado por comas y si no hubiese, ignorarlas y no incluirlas"

    # Guardar el archivo PDF subido
    with open(pdf_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Extraer datos sensibles usando la API
    texts_to_anonymize = extract_sensitive_data_from_pdf(pdf_path, consulta_datos_sensibles)

    # Extraer resolución y número de expediente
    resolution_data = extract_resolution_and_case_number(pdf_path, consulta_resolucion)
    resolution = resolution_data[0] if resolution_data else "sin_resolucion"
    case_number = resolution_data[1] if len(resolution_data) > 1 else "sin_expediente"

    # Crear el nombre del archivo de salida
    sanitized_resolution = sanitize_filename(resolution)
    sanitized_case_number = sanitize_filename(case_number)
    #output_pdf_name = f"Resolucion {sanitized_resolution}_{sanitized_case_number}.pdf"
    output_pdf_name = f"Resolucion {sanitized_resolution}.pdf"

    # Anonimizar el PDF con los datos extraídos
    output_pdf_stream = anonymize_pdf(pdf_path, texts_to_anonymize)

    # Modificar el archivo Word y preparar para descarga
    output_doc_stream = modify_word_file(resolution, case_number)
    output_doc_name = f"FORMATO SAIP {sanitized_case_number}.docx"


    # Descargar el PDF anonimizado
    st.download_button("Descargar PDF Anonimizado", output_pdf_stream, file_name=output_pdf_name)

    # Descargar el documento Word modificado
    if output_doc_stream:
        st.download_button("Descargar Formato SAIP", output_doc_stream, file_name=output_doc_name)

    # Descargar el documento Word modificado (Informe de SAIP Atención Parcial) si hay un solicitante
    if nombre_solicitante:
        output_saip_report_stream = modify_saip_report(resolution, nombre_solicitante)
        output_saip_report_name = f"Informe de SAIP_Atención Parcial_{sanitized_case_number}.docx"

        if output_saip_report_stream:
           st.download_button("Descargar Informe de SAIP Atención Parcial", output_saip_report_stream,
                                   file_name=output_saip_report_name)

    st.success(f"El PDF ha sido anonimizado y está listo para descargar.")




